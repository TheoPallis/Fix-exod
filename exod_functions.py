import os
import re
import docx
from odf.opendocument import load
from odf.text import P
from odf.element import Element
from docx.opc.exceptions import PackageNotFoundError
from codes.exod_ds import reg_amount,production_dir, non_docs,errors_dict
from collections import defaultdict
#pip install odfpy
#pip install python-docx

def clean_amount(amount):
    amount = amount[-200:]
    start_amount_string = amount.find('σε')
    end_amount_string = amount.find('καθίσταται')
    cleaned_amount = amount[start_amount_string + 2:end_amount_string].strip()
    return cleaned_amount

def extract_amount(current_doc_string, reg_amount):
    amount_string = re.findall(reg_amount, current_doc_string)
    if amount_string :
        return clean_amount(amount_string)
    else :
        return "None"


def append_amount_to_dict(r,file,) :
    base_r = os.path.basename(r)
    current_doc = docx.Document(os.path.join(r, file))
    paragraph_text_list = [p.text for p in list(current_doc.paragraphs)]               
    current_doc_string = ''.join(str(x) for x in  paragraph_text_list)     
    amount = extract_amount(current_doc_string,reg_amount)
    # print(f"                         {amount}")
    # amounts_dict[main_folder][0].append(amount)
    # amounts_dict[main_folder][1].append(base_r)

     
# Functions to merge, indent, find non word files

def indent_item(item, main_folder, num=None):
    base_item = os.path.basename(item)
    rel_path_item = os.path.relpath(item, main_folder)
    depth_item = rel_path_item.count(os.sep)
    if num is not None:
        print(f"{ '     ' * depth_item} {num+1}. {base_item}")
    else:
        print(f"{ '     ' * depth_item} {base_item}")
    return base_item, depth_item

def read_document(r, file):
    if file.endswith(".docx") or file.endswith(".doc"):
        return docx.Document(os.path.join(r, file))
    elif file.endswith(".odt") or file.endswith(".odf"):
        return load(os.path.join(r, file))

def find_non_word_document_types(file, non_docs):
    extension = str(file).split('.')[1]
    non_docs[extension].append(file)

def catch_errors(e,r,file,errors_dict,main_folder) :
    base_r = os.path.basename(r)
    full_file = os.path.join(r, file)
    error_message = (str(e), full_file)
    errors_dict[main_folder][0].append(base_r)
    errors_dict[main_folder][1].append(file)
    errors_dict[main_folder][2].append(error_message)
    errors_dict[main_folder][3].append(full_file)
    
def merged_document(r, file, merged_doc):
    if file.endswith(".docx") or file.endswith(".doc"):
        doc = docx.Document(os.path.join(r, file))
        print(f"                Readed doc for {file} done")
        for i,paragraph in enumerate(doc.paragraphs):
            merged_doc.add_paragraph(paragraph.text)
            print(f"                    Paragraph {i+1} added")
    elif file.endswith(".odt") or file.endswith(".odf"):
        doc = load(os.path.join(r, file))
        for paragraph in doc.getElementsByType(P):
            merged_doc.add_paragraph(paragraph.text)  
    return merged_doc
   
def read_all_files(files,r,main_folder,merged_doc,amounts_dict) :
    for k, file in enumerate(files):
        if file.endswith((".docx", ".doc", ".odt", ".odf")):
            try:
                indent_item(file, r, k)
                # Merge to total
                merged_document(r, file, merged_doc)
                print (f"           Merged document done")
                # doc = read_document(r, file)
                # Append amount
                append_amount_to_dict(r,file,)
            except Exception as e:
                 catch_errors(e,r,file,errors_dict,main_folder) 
        else :
            find_non_word_document_types(file, non_docs)

    return errors_dict,amounts_dict[main_folder],non_docs

def merge_all_files_in_filtered_folder(folders_to_search,filtered_folder=None):
    for i, main_folder in enumerate(folders_to_search):
        if filtered_folder is None or main_folder == filtered_folder:
            merged_doc = docx.Document()
            amounts_dict = defaultdict(lambda: ([], []))
            full_path_main_folder = os.path.join(production_dir, main_folder)
            print(f"{(i + 1)} {main_folder}")
            for j, (r, s, files) in enumerate(os.walk(full_path_main_folder)):
                errors_dict, amounts_dict[main_folder], non_docs = read_all_files(files, r, main_folder, merged_doc, amounts_dict)
            merged_doc.save(f"{main_folder}.docx")

    # return errors_dict, amounts_dict, non_docs